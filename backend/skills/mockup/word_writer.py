import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from parser import ScreenSpec
from image_writer import write_mockup_image


_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = _WHITE
        _shade_cell(cell, "1F4E79")
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
    doc.add_paragraph()


def write_word(spec: ScreenSpec, descriptions: dict, output_path: Path):
    doc = Document()

    doc.add_heading(spec.screen_title, level=1)

    # Mockup image
    img_bytes = write_mockup_image(spec)
    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.3))
    doc.add_paragraph()

    doc.add_heading("Descriere Generală", 2)
    doc.add_paragraph(descriptions.get("descriere_generala", ""))

    doc.add_heading("Reguli Business", 2)
    for rule in descriptions.get("reguli_business", []):
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("Zone Ecran", 2)
    for sec in spec.sections:
        doc.add_heading(sec.title, 3)

        if sec.filter_fields:
            doc.add_heading("Zona de Filtrare", 4)
            rows = [
                [
                    f.label,
                    "Da" if f.mandatory else "Nu",
                    descriptions.get("descrieri_filtre", {}).get(f.label, ""),
                ]
                for f in sec.filter_fields
            ]
            _add_table(doc, ["Câmp", "Obligatoriu", "Descriere"], rows)

        if sec.buttons:
            doc.add_heading("Butoane și Acțiuni", 4)
            rows = [
                [
                    b.label,
                    "Direct" if b.group == "direct" else "Acțiune",
                    descriptions.get("descrieri_butoane", {}).get(b.label, ""),
                ]
                for b in sec.buttons
            ]
            _add_table(doc, ["Buton", "Tip", "Descriere"], rows)

        if sec.columns:
            doc.add_heading("Coloane Grid", 4)
            rows = [
                [
                    c.name,
                    descriptions.get("descrieri_coloane", {}).get(c.name, ""),
                ]
                for c in sec.columns
            ]
            _add_table(doc, ["Coloană", "Descriere"], rows)

    doc.add_heading("Mod de Lucru", 2)
    for step in descriptions.get("mod_de_lucru", []):
        doc.add_paragraph(step, style="List Number")

    doc.save(str(output_path))
