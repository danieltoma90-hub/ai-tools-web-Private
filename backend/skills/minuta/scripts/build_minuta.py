"""Generator minută F.05 din JSON intermediar — design TotalSoft (PV style).

Pornește de la `F05_minuta_template.docx` care conține:
  - header real cu logo TotalSoft + linie albastră
  - footer real cu logo Charisma + text central + paginare automată
  - bloc titlu (MINUTA INTALNIRII + subtitlu + linie roșie)
  - tabel de identificare cu etichete pe fundal albastru deschis

și adaugă conținutul variabil (secțiuni cu heading-uri subliniate, tabele de
clarificări cu header navy, pași următori) folosind paleta vizuală TotalSoft.

Utilizare:
    python build_minuta.py <input.json> <output.docx>

Format JSON: vezi `examples/json_intermediate.json`.
"""
from __future__ import annotations

import argparse
import io
import json
import re
import sys

# Fix Romanian diacritics on Windows consoles (cp1252 → UTF-8)
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() not in ("utf-8", "utf8"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

# ─────────────────────────────────────────────────────────────────────────
# Paleta vizuală TotalSoft (extrasă din PV Postimplementare)
# ─────────────────────────────────────────────────────────────────────────

NAVY = "1F3864"             # titlu, heading-uri, header tabele, etichete
BLUE_MID = "2E5496"         # subtitlu, subheading-uri
BLUE_LIGHT_FILL = "F2F6FB"  # fundal etichete + rânduri alternante (zebra)
RED = "C00000"              # linii accent
GREY_TEXT = "595959"        # text footer / secundar
BORDER_GREY = "BFBFBF"      # chenar tabele
WHITE = "FFFFFF"

CONTENT_TABLE_WIDTH = 9638  # lățime utilă (margini 2cm)

# ─────────────────────────────────────────────────────────────────────────
# Helpers XML de stil
# ─────────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = BORDER_GREY, size: str = "4") -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_margins(cell, top=60, left=110, bottom=60, right=110) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)


def _set_cell_valign(cell, val: str = "center") -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tc_pr.append(v)


def _set_run_lang(run, lang: str = "ro-RO") -> None:
    rpr = run._element.get_or_add_rPr()
    existing = rpr.find(qn("w:lang"))
    if existing is not None:
        rpr.remove(existing)
    lang_el = OxmlElement("w:lang")
    lang_el.set(qn("w:val"), lang)
    rpr.append(lang_el)


def _add_bottom_border(paragraph, color: str = NAVY, size: str = "10", space: str = "4") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.insert(0, p_bdr)


def _keep_next(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        kn = OxmlElement("w:keepNext")
        p_pr.insert(0, kn)


# ─────────────────────────────────────────────────────────────────────────
# Markup inline **bold**
# ─────────────────────────────────────────────────────────────────────────

def _add_runs_with_markup(paragraph, text: str) -> None:
    """Adaugă runs interpretând **bold**. Suportă \n → paragrafe noi."""
    lines = text.split("\n")
    parent = paragraph._element.getparent()
    p_pr = paragraph._element.find(qn("w:pPr"))
    current_para = paragraph

    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            new_p = OxmlElement("w:p")
            if p_pr is not None:
                new_p.append(deepcopy(p_pr))
            idx = list(parent).index(current_para._element)
            parent.insert(idx + 1, new_p)
            current_para = Paragraph(new_p, paragraph._parent)

        parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = current_para.add_run(part[2:-2])
                run.bold = True
            else:
                run = current_para.add_run(part)
            _set_run_lang(run)


def _add_runs_with_markup_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        _set_run_lang(run)


# ─────────────────────────────────────────────────────────────────────────
# Placeholders
# ─────────────────────────────────────────────────────────────────────────

def _format_participanti(participanti: dict[str, list[str]]) -> str:
    lines = []
    for grup, persoane in participanti.items():
        if persoane:
            lines.append(f"{grup}: {', '.join(persoane)}")
    return "\n".join(lines)


def _replace_placeholder_in_paragraph(para, mapping: dict[str, str]) -> None:
    full_text = "".join(r.text for r in para.runs)
    matched_ph = None
    for ph in mapping:
        if ph in full_text:
            matched_ph = ph
            break
    if matched_ph is None:
        return

    new_text = full_text.replace(matched_ph, mapping[matched_ph])

    saved_rpr = None
    if para.runs:
        rpr = para.runs[0]._element.find(qn("w:rPr"))
        if rpr is not None:
            saved_rpr = deepcopy(rpr)

    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    lines = new_text.split("\n")
    run = para.add_run(lines[0])
    if saved_rpr is not None:
        run._element.insert(0, deepcopy(saved_rpr))

    parent = para._element.getparent()
    p_pr = para._element.find(qn("w:pPr"))
    for i, line in enumerate(lines[1:], start=1):
        new_p = OxmlElement("w:p")
        if p_pr is not None:
            new_p.append(deepcopy(p_pr))
        new_r = OxmlElement("w:r")
        if saved_rpr is not None:
            new_r.append(deepcopy(saved_rpr))
        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        new_p.append(new_r)
        idx = list(parent).index(para._element)
        parent.insert(idx + i, new_p)


def _replace_all_placeholders(doc, mapping: dict[str, str]) -> None:
    for para in doc.paragraphs:
        _replace_placeholder_in_paragraph(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_placeholder_in_paragraph(para, mapping)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                _replace_placeholder_in_paragraph(para, mapping)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_placeholder_in_paragraph(para, mapping)


# ─────────────────────────────────────────────────────────────────────────
# Elemente de conținut
# ─────────────────────────────────────────────────────────────────────────

def _add_section_heading(doc, number: int, title: str) -> None:
    p = doc.add_paragraph()
    _keep_next(p)
    _add_bottom_border(p, NAVY, "10", "4")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    _set_run_lang(run)


def _add_subheading(doc, text: str) -> None:
    p = doc.add_paragraph()
    _keep_next(p)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor.from_string(BLUE_MID)
    _set_run_lang(run)


def _add_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_runs_with_markup(p, text)


def _add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("•\t")
    _set_run_lang(run)
    _add_runs_with_markup_inline(p, text)


def _add_numbered(doc, text: str, number: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{number}.\t")
    _set_run_lang(run)
    _add_runs_with_markup_inline(p, text)


def _add_content_table(doc, header: list[str], rows: list[list[str]],
                       col_widths: list[int] | None = None) -> None:
    """Tabel cu header navy + text alb și rânduri zebra (2/3/4 coloane)."""
    if not rows:
        return

    ncols = len(header)
    table = doc.add_table(rows=len(rows) + 1, cols=ncols)
    table.autofit = False

    tbl_pr = table._element.find(qn("w:tblPr"))
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    if col_widths is None:
        if ncols == 2:
            col_widths = [3000, 6638]
        elif ncols == 3:
            col_widths = [600, 4000, 5038]
        elif ncols == 4:
            col_widths = [600, 3400, 2200, 3438]
        else:
            unit = CONTENT_TABLE_WIDTH // ncols
            col_widths = [unit] * ncols

    tbl_grid = table._element.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        table._element.remove(tbl_grid)
    new_grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        new_grid.append(gc)
    tbl_pr.addnext(new_grid)

    hdr_row = table.rows[0]
    tr_pr = hdr_row._element.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))

    for i, h in enumerate(header):
        cell = hdr_row.cells[i]
        _set_cell_shading(cell, NAVY)
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        _set_cell_valign(cell, "center")
        para = cell.paragraphs[0]
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(WHITE)
        _set_run_lang(run)

    for row_idx, row_data in enumerate(rows, start=1):
        is_even = (row_idx % 2 == 0)
        for col_idx in range(ncols):
            val = row_data[col_idx] if col_idx < len(row_data) else ""
            cell = table.rows[row_idx].cells[col_idx]
            if is_even:
                _set_cell_shading(cell, BLUE_LIGHT_FILL)
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            _set_cell_valign(cell, "center")
            para = cell.paragraphs[0]
            for r in list(para.runs):
                r._element.getparent().remove(r._element)
            _add_runs_with_markup(para, val)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_signature_block(doc) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_margins(cell)
    h_cell = table.rows[0].cells[0]
    _set_cell_shading(h_cell, BLUE_LIGHT_FILL)
    run = h_cell.paragraphs[0].add_run("Reprezentant TotalSoft")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    sig_cell = table.rows[2].cells[0]
    run = sig_cell.paragraphs[0].add_run("Semnatura")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(GREY_TEXT)


# ─────────────────────────────────────────────────────────────────────────
# Procesare blocuri
# ─────────────────────────────────────────────────────────────────────────

def _add_section_content(doc, blocks: list[dict]) -> None:
    for block in blocks:
        btype = block.get("type")
        if btype == "subheading":
            _add_subheading(doc, block["text"])
        elif btype == "paragraph":
            _add_paragraph(doc, block["text"])
        elif btype == "bullets":
            for item in block.get("items", []):
                _add_bullet(doc, item)
        elif btype in ("table_2col", "table"):
            header = block.get("header", ["Parametru", "Explicatie agreata"])
            rows = block.get("rows", [])
            col_widths = block.get("col_widths")
            _add_content_table(doc, header, rows, col_widths)
        elif btype == "image":
            caption = block.get("caption", "")
            if caption:
                _add_subheading(doc, caption)
            _add_paragraph(doc, f"[Imagine: {block.get('path', '')}]")
        else:
            raise ValueError(f"Tip bloc necunoscut: {btype}")


def _add_next_steps(doc, steps: list[dict], section_number: int) -> None:
    _add_section_heading(doc, section_number, "Pași următori")
    for idx, step in enumerate(steps, start=1):
        resp = step.get("responsabil", "").strip()
        action = step.get("actiune", "").strip()
        text = f"{resp}: {action}" if resp else action
        _add_numbered(doc, text, idx)


# ─────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────

def _build_footer_text(meta: dict) -> str:
    client = meta.get("nume_client", "")
    return f"Minuta Intalnirii | {client}" if client else "Minuta Intalnirii"


def build_minuta(json_path: Path, output_path: Path, template_path: Path | None = None) -> None:
    if template_path is None:
        template_path = Path(__file__).parent.parent / "template" / "F05_minuta_template.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Template lipsă: {template_path}")

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    meta = data.get("meta", {})
    doc = Document(str(template_path))

    cod_formular = meta.get("cod_formular", "F.05 – Minuta Intalnirii")
    participanti_text = _format_participanti(meta.get("participanti", {}))
    nume_client = meta.get("nume_client", "")
    subiect = meta.get("subiect", "")

    mapping = {
        "{{SUBTITLU}}": subiect or cod_formular,
        "{{TITLU_PROIECT}}": nume_client,
        "{{COD_PROIECT}}": meta.get("cod_proiect", ""),
        "{{NUME_CLIENT}}": nume_client,
        "{{NUMAR_CONTRACT}}": meta.get("numar_contract", ""),
        "{{SUBIECT}}": subiect,
        "{{INITIATOR}}": meta.get("initiator", ""),
        "{{PARTICIPANTI}}": participanti_text,
        "{{DISTRIBUIT}}": meta.get("distribuit", ""),
        "{{LOCATIA}}": meta.get("locatia", ""),
        "{{DATA}}": meta.get("data", ""),
        "{{DURATA}}": meta.get("durata", ""),
        "{{FOOTER_TEXT}}": _build_footer_text(meta),
    }
    _replace_all_placeholders(doc, mapping)

    section_idx = 1
    context = data.get("context_si_scop")
    if context:
        _add_section_heading(doc, section_idx, "Context și scop")
        _add_section_content(doc, context)
        section_idx += 1

    for sec in data.get("sectiuni", []):
        _add_section_heading(doc, section_idx, sec["titlu"])
        _add_section_content(doc, sec.get("blocuri", []))
        section_idx += 1

    pasi = data.get("pasi_urmatori", [])
    if pasi:
        _add_next_steps(doc, pasi, section_idx)
        section_idx += 1

    if data.get("include_signature"):
        _add_signature_block(doc)

    doc.save(str(output_path))
    print(f"Minuta generată: {output_path}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generează o minută F.05 (design TotalSoft).")
    parser.add_argument("json_input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--template", type=Path, default=None)
    args = parser.parse_args()
    try:
        build_minuta(args.json_input, args.output, args.template)
        return 0
    except Exception as e:
        print(f"Eroare: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
