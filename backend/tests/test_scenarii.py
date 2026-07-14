import io
import tempfile
from pathlib import Path as _P
from unittest.mock import patch

from docx import Document

from main import app
from auth import verify_token


def _spec_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Modul Test", level=1)
    doc.add_heading("Capitol Unu", level=2)
    doc.add_paragraph("Descriere functionalitate.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _spec_tempfile() -> str:
    fd, name = tempfile.mkstemp(suffix=".docx")
    import os as _os
    _os.close(fd)
    _P(name).write_bytes(_spec_bytes())
    return name


async def test_estimate_without_auth_returns_401(client):
    response = await client.post(
        "/api/scenarii/estimate",
        json={"storage_path": "scenarii/x.docx", "filename": "spec.docx"},
    )
    assert response.status_code == 401


async def test_estimate_wrong_format_returns_422(client):
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        response = await client.post(
            "/api/scenarii/estimate",
            json={"storage_path": "scenarii/x.pdf", "filename": "spec.pdf"},
            headers={"Authorization": "Bearer fake"},
        )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 422


async def test_estimate_missing_upload_returns_422(client):
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        with patch("routers.scenarii.download_upload", side_effect=Exception("object not found")):
            response = await client.post(
                "/api/scenarii/estimate",
                json={"storage_path": "scenarii/lipsa.docx", "filename": "spec.docx"},
                headers={"Authorization": "Bearer fake"},
            )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 422
    assert "nu a fost găsit" in response.json()["detail"]


async def test_estimate_then_generate_then_job_done(client, monkeypatch, tmp_path):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test")

    out_xlsx = tmp_path / "out.xlsx"
    out_xlsx.write_bytes(b"PK-fake-xlsx")

    async def fake_pipeline(input_path, engine="claude", on_step=None):
        return out_xlsx, {"core_count": 100, "specific_count": 2,
                          "excluded_count": 5, "requirements": 2}

    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        with patch("routers.scenarii.download_upload", return_value=_P(_spec_tempfile())):
            est_res = await client.post(
                "/api/scenarii/estimate",
                json={"storage_path": "scenarii/x.docx", "filename": "spec.docx"},
                headers={"Authorization": "Bearer fake"},
            )
        assert est_res.status_code == 200
        est = est_res.json()
        assert "estimate_id" in est
        assert "est_minutes_free" in est

        with patch("routers.scenarii.run_scenarii_pipeline", side_effect=fake_pipeline), \
             patch("routers.scenarii.upload_file", return_value="scenarii/x.xlsx"):
            gen_res = await client.post(
                "/api/scenarii/generate",
                json={"estimate_id": est["estimate_id"], "engine": "claude"},
                headers={"Authorization": "Bearer fake"},
            )
        assert gen_res.status_code == 200
        job_id = gen_res.json()["job_id"]

        job_res = await client.get(
            f"/api/scenarii/job/{job_id}",
            headers={"Authorization": "Bearer fake"},
        )
        assert job_res.status_code == 200
        job = job_res.json()
        assert job["status"] == "done"
        assert job["engine"] == "claude"
        assert job["filename"].endswith(".xlsx")
        assert job["summary"]["specific_count"] == 2
        assert job["xlsx_b64"]
    finally:
        app.dependency_overrides.clear()


async def test_generate_with_expired_estimate_returns_404(client):
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        response = await client.post(
            "/api/scenarii/generate",
            json={"estimate_id": "inexistent", "engine": "claude"},
            headers={"Authorization": "Bearer fake"},
        )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 404


async def test_generate_invalid_engine_returns_422(client):
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        response = await client.post(
            "/api/scenarii/generate",
            json={"estimate_id": "x", "engine": "mistral"},
            headers={"Authorization": "Bearer fake"},
        )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 422


async def test_generate_missing_api_key_returns_500(client, monkeypatch):
    monkeypatch.delenv("GROQ_API_KEY", raising=False)
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        with patch("routers.scenarii.download_upload", return_value=_P(_spec_tempfile())):
            est_res = await client.post(
                "/api/scenarii/estimate",
                json={"storage_path": "scenarii/x.docx", "filename": "spec.docx"},
                headers={"Authorization": "Bearer fake"},
            )
        est = est_res.json()
        response = await client.post(
            "/api/scenarii/generate",
            json={"estimate_id": est["estimate_id"], "engine": "groq"},
            headers={"Authorization": "Bearer fake"},
        )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 500
    assert "GROQ_API_KEY" in response.json()["detail"]


async def test_job_not_found_returns_404(client):
    app.dependency_overrides[verify_token] = lambda: {"id": "u1"}
    try:
        response = await client.get(
            "/api/scenarii/job/nu-exista",
            headers={"Authorization": "Bearer fake"},
        )
    finally:
        app.dependency_overrides.clear()
    assert response.status_code == 404
