import pytest
from unittest.mock import patch
from pathlib import Path
from pipelines.minuta_pipeline import extract_vtt_text, extract_docx_text, run_minuta_pipeline


def test_extract_vtt_text_returns_plain_text(tmp_path):
    vtt = tmp_path / "transcript.vtt"
    vtt.write_text(
        "WEBVTT\n\n00:00:01.000 --> 00:00:03.000\nIon Popescu: Buna ziua.\n\n"
        "00:00:04.000 --> 00:00:07.000\nMaria Ionescu: Multumim ca ati venit.",
        encoding="utf-8"
    )
    text = extract_vtt_text(vtt)
    assert "Ion Popescu: Buna ziua." in text
    assert "WEBVTT" not in text
    assert "00:00:01" not in text


def test_extract_docx_text_returns_paragraphs(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Ion Popescu: Buna ziua tuturor.")
    doc.add_paragraph("Maria Ionescu: Incepem sedinta.")
    docx_path = tmp_path / "transcript.docx"
    doc.save(str(docx_path))

    text = extract_docx_text(docx_path)
    assert "Ion Popescu: Buna ziua tuturor." in text
    assert "Maria Ionescu: Incepem sedinta." in text


@pytest.mark.asyncio
async def test_run_minuta_pipeline_calls_claude_and_returns_docx(tmp_path):
    vtt = tmp_path / "transcript.vtt"
    vtt.write_text("WEBVTT\n\n00:00:01.000 --> 00:00:05.000\nDaniel: Test sedinta.", encoding="utf-8")

    mock_meta = {
        "cod_formular": "F.05",
        "subiect": "Test Sedinta",
        "data": "25.06.2026",
        "durata": "30min",
        "locatia": "Microsoft Teams",
        "initiator": "Daniel Toma",
        "cod_proiect": "TEST",
        "numar_contract": "",
        "distribuit": "",
        "nume_client": "Client Test",
        "participanti": {"TotalSoft": ["Daniel Toma"], "Client": []}
    }
    mock_sections = {"context_si_scop": None, "sectiuni": []}
    mock_actions = []

    with patch("pipelines.minuta_pipeline.extract_metadata", return_value=mock_meta), \
         patch("pipelines.minuta_pipeline.extract_sections", return_value=mock_sections), \
         patch("pipelines.minuta_pipeline.extract_action_items", return_value=mock_actions):
        docx_path, preview_html = await run_minuta_pipeline(vtt, api_key="fake-key")

    assert docx_path.exists()
    assert docx_path.suffix == ".docx"
    assert isinstance(preview_html, str)
    assert "Test Sedinta" in preview_html
    docx_path.unlink()
