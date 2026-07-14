# -*- coding: utf-8 -*-
"""Teste pentru pipeline-ul scenarii (catalog standard + cerinte specifice)."""
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from skills.scenarii import ai_gen
from skills.scenarii.spec_parser import parse_client_spec
from pipelines.scenarii_pipeline import (
    _apply_exclusions,
    _load_catalog,
    estimate_scenarii_job,
    run_scenarii_pipeline,
)


def _make_spec(tmp_path: Path, with_specific: bool = True) -> Path:
    """Docx sintetic in structura Model Company, cu o sectiune de cerinte specifice."""
    doc = Document()
    doc.add_heading("Introducere", level=1)
    doc.add_paragraph("Despre document.")
    doc.add_heading("Nomenclatoare", level=1)
    doc.add_heading("Articole", level=2)
    doc.add_paragraph("Descriere articole.")
    doc.add_heading("Campuri obligatorii", level=3)
    doc.add_paragraph("Campurile x, y.")
    if with_specific:
        doc.add_heading("Cerinte specifice identificate in urma Analizei", level=3)
        doc.add_heading("Codificare automata articole CR.00.01", level=4)
        doc.add_paragraph("Codul articolului se genereaza automat dupa masca GRUPA-NNNNN.")
        doc.add_heading("Alerta stoc minim", level=4)
        doc.add_paragraph("Sistemul trimite notificare cand stocul scade sub minim.")
    doc.add_heading("Fluxuri de business", level=1)
    doc.add_heading("Modul Achizitii", level=2)
    doc.add_heading("Achizitii intern – Factura de aprovizionare", level=3)
    doc.add_paragraph("Fluxul standard de facturare.")
    path = tmp_path / "spec.docx"
    doc.save(str(path))
    return path


def test_catalog_loads_and_has_scenarios():
    catalog = _load_catalog()
    total = sum(len(v) for v in catalog.values())
    assert total >= 150
    ids = [s["ID"] for items in catalog.values() for s in items]
    assert len(ids) == len(set(ids)), "ID-uri duplicate in catalog"


def test_parser_extracts_specific_requirements(tmp_path):
    spec = parse_client_spec(_make_spec(tmp_path))
    assert len(spec.requirements) == 2
    r = spec.requirements[0]
    assert r.code == "CR.00.01"
    assert "Codificare automata" in r.title
    assert r.module_h2 == "Articole"
    assert "masca" in r.text


def test_exclusions_remove_missing_chapters(tmp_path):
    spec = parse_client_spec(_make_spec(tmp_path))
    catalog = _load_catalog()
    before = sum(len(v) for v in catalog.values())
    excluded = _apply_exclusions(catalog, spec.titles)
    after = sum(len(v) for v in catalog.values())
    assert before == after + len(excluded)
    # spec-ul sintetic nu are Vanzari/Financiar etc. -> excluderi masive
    assert len(excluded) > 20
    # scenariile de nucleu (fara std_ref) raman mereu
    assert any(s["ID"].startswith("CFG-") for v in catalog.values() for s in v)
    # dependentele catre scenariile excluse au fost curatate
    excluded_ids = {s["ID"] for _, s in excluded}
    for items in catalog.values():
        for s in items:
            assert not (set(s.get("deps", [])) & excluded_ids)


def test_estimate_returns_both_variants(tmp_path):
    est = estimate_scenarii_job(_make_spec(tmp_path))
    assert est["requirements"] == 2
    assert est["est_minutes"] >= 1
    assert est["est_minutes_free"] >= 1


@pytest.mark.asyncio
async def test_pipeline_with_mocked_ai(tmp_path, monkeypatch):
    async def fake_generate(reqs, engine, on_step=None):
        return [{
            "cerinta_index": i,
            "sheet": "🏷️ Articole",
            "Scenariu": f"Test cerinta {i}",
            "Obiectiv": "Obiectiv",
            "Pasi de Test": "1. Pas",
            "Rezultat Asteptat": "OK",
            "Tip Test": "Functional - Pozitiv",
            "Prioritate": "High",
        } for i in range(len(reqs))]

    async def fake_deps(core_index, new_index, engine, on_step=None):
        first_core = core_index[0][0]
        return {new_index[0][0]: [first_core]}

    monkeypatch.setattr(ai_gen, "generate_specific_scenarios", fake_generate)
    monkeypatch.setattr(ai_gen, "generate_dependencies", fake_deps)

    xlsx, summary = await run_scenarii_pipeline(_make_spec(tmp_path), engine="claude")
    try:
        assert summary["specific_count"] == 2
        assert summary["excluded_count"] > 0

        import io
        wb = load_workbook(io.BytesIO(xlsx.read_bytes()))  # fara handle pe fisier (Windows)
        assert "🚫 Excluse vs Standard" in wb.sheetnames
        assert "🗺️ Plan Executie" in wb.sheetnames
        ws = wb["🏷️ Articole"]
        cs_ids = [row[0] for row in ws.iter_rows(min_row=2, values_only=True)
                  if row[0] and "-CS-" in str(row[0])]
        assert len(cs_ids) == 2
        # scenariul specific are codul cerintei in Observatii
        obs = [row[12] for row in ws.iter_rows(min_row=2, values_only=True)
               if row[0] in cs_ids]
        assert any("CR.00.01" in str(o) for o in obs)
    finally:
        xlsx.unlink(missing_ok=True)


def test_ai_gen_batching():
    from skills.scenarii.spec_parser import SpecificRequirement
    reqs = [
        SpecificRequirement(code="", title=f"R{i}", text="x" * 4000,
                            module_h2="Articole", module_no="4.1",
                            chapter_no="4.1.9", item_no=f"4.1.9.{i}")
        for i in range(6)
    ]
    batches = ai_gen.batch_requirements(reqs, 9_000)
    assert sum(len(b) for b in batches) == 6
    assert all(len(b) <= 2 for b in batches)  # 2 x 4000 < 9000, 3 x 4000 > 9000


def test_parse_json_block_variants():
    assert ai_gen.parse_json_block('{"a": 1}') == {"a": 1}
    assert ai_gen.parse_json_block('```json\n{"a": 1}\n```') == {"a": 1}
    assert ai_gen.parse_json_block('text inainte {"a": 1} dupa') == {"a": 1}
